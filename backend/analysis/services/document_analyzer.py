import os
import magic
import PyPDF2
from docx import Document as DocxDocument
import pandas as pd
from datetime import datetime

try:
    from PIL import Image
    PILLOW_AVAILABLE = True
except ImportError:
    PILLOW_AVAILABLE = False
    print("Warning: Pillow not installed. Image analysis will be limited.")

class DocumentAnalyzer:
    @staticmethod
    def analyze(file_path, original_name, mime_type):
        file_stats = os.stat(file_path)
        analysis = {
            'basic_info': {
                'filename': original_name,
                'size_bytes': file_stats.st_size,
                'size_mb': round(file_stats.st_size / (1024 * 1024), 2),
                'created': datetime.fromtimestamp(file_stats.st_ctime).isoformat(),
                'modified': datetime.fromtimestamp(file_stats.st_mtime).isoformat(),
            },
            'mime_type': mime_type,
            'analysis_timestamp': datetime.now().isoformat(),
        }
        
        if mime_type == 'application/pdf':
            analysis.update(DocumentAnalyzer._analyze_pdf(file_path))
        elif mime_type.startswith('image/'):
            analysis.update(DocumentAnalyzer._analyze_image(file_path))
        elif 'document' in mime_type or mime_type == 'application/msword':
            analysis.update(DocumentAnalyzer._analyze_word(file_path))
        elif 'spreadsheet' in mime_type or 'excel' in mime_type:
            analysis.update(DocumentAnalyzer._analyze_spreadsheet(file_path))
        elif mime_type == 'text/plain':
            analysis.update(DocumentAnalyzer._analyze_text(file_path))
        elif mime_type == 'text/csv':
            analysis.update(DocumentAnalyzer._analyze_csv(file_path))
        
        return analysis
    
    @staticmethod
    def _analyze_pdf(file_path):
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                return {
                    'document_type': 'pdf',
                    'pages': len(reader.pages),
                    'is_encrypted': reader.is_encrypted,
                    'metadata': reader.metadata if reader.metadata else {},
                }
        except Exception as e:
            return {'error': f'PDF analysis failed: {str(e)}'}
    
    @staticmethod
    def _analyze_image(file_path):
        if not PILLOW_AVAILABLE:
            return {
                'document_type': 'image',
                'note': 'Pillow not installed - detailed image analysis unavailable',
            }
        try:
            with Image.open(file_path) as img:
                return {
                    'document_type': 'image',
                    'width': img.width,
                    'height': img.height,
                    'format': img.format,
                    'mode': img.mode,
                }
        except Exception as e:
            return {'error': f'Image analysis failed: {str(e)}'}
    
    @staticmethod
    def _analyze_word(file_path):
        try:
            doc = DocxDocument(file_path)
            return {
                'document_type': 'word',
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'sections': len(doc.sections),
            }
        except Exception as e:
            return {'error': f'Word analysis failed: {str(e)}'}
    
    @staticmethod
    def _analyze_spreadsheet(file_path):
        try:
            df = pd.read_excel(file_path, sheet_name=None)
            sheets_info = {}
            for sheet_name, data in df.items():
                sheets_info[sheet_name] = {
                    'rows': len(data),
                    'columns': len(data.columns),
                    'column_names': list(data.columns[:5]),
                }
            return {
                'document_type': 'spreadsheet',
                'sheet_count': len(df),
                'sheets': sheets_info,
            }
        except Exception as e:
            return {'error': f'Spreadsheet analysis failed: {str(e)}'}
    
    @staticmethod
    def _analyze_text(file_path):
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
                return {
                    'document_type': 'text',
                    'character_count': len(content),
                    'word_count': len(content.split()),
                    'line_count': content.count('\n') + 1,
                }
        except Exception as e:
            return {'error': f'Text analysis failed: {str(e)}'}
    
    @staticmethod
    def _analyze_csv(file_path):
        try:
            df = pd.read_csv(file_path)
            return {
                'document_type': 'csv',
                'rows': len(df),
                'columns': len(df.columns),
                'column_names': list(df.columns),
            }
        except Exception as e:
            return {'error': f'CSV analysis failed: {str(e)}'}